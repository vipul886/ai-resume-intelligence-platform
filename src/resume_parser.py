"""
Resume Parser Module
Handles PDF/DOCX parsing and text extraction from resumes.
"""

import re
import PyPDF2
import pdfplumber
from typing import Dict, Optional, List
from pathlib import Path
import logging
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Advanced resume parser that extracts text and metadata from PDF and DOCX files.
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
        logger.info("ResumeParser initialized")
    
    def parse_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using multiple methods for reliability.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        try:
            # Try pdfplumber first (better for complex layouts)
            text = self._parse_with_pdfplumber(file_path)
            if text and len(text.strip()) > 50:
                return text
            
            # Fallback to PyPDF2
            text = self._parse_with_pypdf2(file_path)
            return text
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    def _parse_with_pdfplumber(self, file_path: str) -> str:
        """Parse PDF using pdfplumber."""
        text_content = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        return "\n".join(text_content)
    
    def _parse_with_pypdf2(self, file_path: str) -> str:
        """Parse PDF using PyPDF2 as fallback."""
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        return "\n".join(text_content)
    
    def parse_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    def parse_file(self, file_path: str) -> Dict[str, any]:
        """
        Parse resume file and extract text with metadata.
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        # Extract text based on file type
        if extension == '.pdf':
            text = self.parse_pdf(str(file_path))
        elif extension == '.docx':
            text = self.parse_docx(str(file_path))
        elif extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported format: {extension}")
        
        # Clean and normalize text
        text = self.clean_text(text)
        
        # Extract basic metadata
        metadata = self.extract_metadata(text)
        
        return {
            "raw_text": text,
            "cleaned_text": text,
            "file_name": file_path.name,
            "file_type": extension,
            "metadata": metadata,
            "word_count": len(text.split()),
            "char_count": len(text)
        }
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.,;:\-\(\)@/]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def extract_metadata(self, text: str) -> Dict[str, any]:
        """
        Extract basic metadata from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary with metadata
        """
        metadata = {}
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        metadata['emails'] = emails[:1] if emails else []
        
        # Extract phone numbers
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        metadata['phones'] = [p[0] + p[1] if isinstance(p, tuple) else p for p in phones[:2]]
        
        # Extract LinkedIn
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin = re.findall(linkedin_pattern, text.lower())
        metadata['linkedin'] = linkedin[0] if linkedin else None
        
        # Extract GitHub
        github_pattern = r'github\.com/[\w-]+'
        github = re.findall(github_pattern, text.lower())
        metadata['github'] = github[0] if github else None
        
        return metadata
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Identify and extract common resume sections.
        
        Args:
            text: Resume text
            
        Returns:
            Dictionary mapping section names to content
        """
        sections = {}
        
        # Common section headers
        section_patterns = {
            'summary': r'(professional summary|summary|profile|objective)',
            'experience': r'(work experience|experience|employment history|professional experience)',
            'education': r'(education|academic background|qualifications)',
            'skills': r'(skills|technical skills|core competencies|expertise)',
            'projects': r'(projects|portfolio)',
            'certifications': r'(certifications|certificates|licenses)'
        }
        
        text_lower = text.lower()
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text_lower)
            if match:
                start_idx = match.start()
                # Find next section or end of text
                next_section_idx = len(text)
                for other_pattern in section_patterns.values():
                    if other_pattern != pattern:
                        next_match = re.search(other_pattern, text_lower[start_idx + 50:])
                        if next_match:
                            candidate_idx = start_idx + 50 + next_match.start()
                            if candidate_idx < next_section_idx:
                                next_section_idx = candidate_idx
                
                sections[section_name] = text[start_idx:next_section_idx].strip()
        
        return sections

